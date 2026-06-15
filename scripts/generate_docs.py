# -*- coding: utf-8 -*-
"""生成东软环保公众监督系统 Word 文档"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, 'docs')

TABLES = [
    {
        'name': 'aqi_level',
        'cn': '空气质量指数级别表',
        'desc': '存储AQI空气质量指数各级别定义',
        'fields': [
            ('id', 'INT', 'PK', '级别ID', '自增主键'),
            ('level_name', 'VARCHAR(20)', '', '级别名称', '如：优、良、轻度污染'),
            ('min_value', 'INT', '', 'AQI最小值', '该级别AQI下限'),
            ('max_value', 'INT', '', 'AQI最大值', '该级别AQI上限'),
            ('color_code', 'VARCHAR(20)', '', '显示颜色', '前端展示用色值'),
            ('health_impact', 'VARCHAR(200)', '', '健康影响说明', '该级别对健康的影响'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'grid_province',
        'cn': '系统网格覆盖省区域表',
        'desc': '存储系统网格覆盖的省级区域信息',
        'fields': [
            ('id', 'INT', 'PK', '省ID', '自增主键'),
            ('province_name', 'VARCHAR(50)', '', '省名称', '如：辽宁省'),
            ('province_code', 'VARCHAR(10)', '', '省编码', '省份简称编码'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'grid_city',
        'cn': '系统网格覆盖市区区域表',
        'desc': '存储系统网格覆盖的市级区域信息',
        'fields': [
            ('id', 'INT', 'PK', '市ID', '自增主键'),
            ('province_id', 'INT', 'FK', '所属省ID', '关联网格省表'),
            ('city_name', 'VARCHAR(50)', '', '市/区名称', '如：沈阳市'),
            ('city_code', 'VARCHAR(10)', '', '市编码', '城市编码'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'sys_admin',
        'cn': '系统管理员表',
        'desc': '存储系统管理员账号信息',
        'fields': [
            ('id', 'INT', 'PK', '管理员ID', '自增主键'),
            ('admin_code', 'VARCHAR(50)', 'UK', '管理员账号', '登录账号，唯一'),
            ('admin_name', 'VARCHAR(50)', '', '管理员姓名', '管理员真实姓名'),
            ('password', 'VARCHAR(100)', '', '登录密码', '登录密码'),
            ('phone', 'VARCHAR(20)', '', '联系电话', '手机号码'),
            ('status', 'TINYINT', '', '状态', '1启用 0禁用'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'supervisor',
        'cn': '公众监督员表',
        'desc': '存储公众监督员账号信息',
        'fields': [
            ('id', 'INT', 'PK', '监督员ID', '自增主键'),
            ('supervisor_code', 'VARCHAR(50)', 'UK', '监督员账号', '登录账号，唯一'),
            ('supervisor_name', 'VARCHAR(50)', '', '监督员姓名', '监督员真实姓名'),
            ('password', 'VARCHAR(100)', '', '登录密码', '登录密码'),
            ('phone', 'VARCHAR(20)', '', '联系电话', '手机号码'),
            ('status', 'TINYINT', '', '状态', '1启用 0禁用'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'grid_member',
        'cn': '空气质量监测网格员表',
        'desc': '存储网格员账号及负责区域信息',
        'fields': [
            ('id', 'INT', 'PK', '网格员ID', '自增主键'),
            ('member_code', 'VARCHAR(50)', 'UK', '网格员账号', '登录账号，唯一'),
            ('member_name', 'VARCHAR(50)', '', '网格员姓名', '网格员真实姓名'),
            ('password', 'VARCHAR(100)', '', '登录密码', '登录密码'),
            ('phone', 'VARCHAR(20)', '', '联系电话', '手机号码'),
            ('province_id', 'INT', 'FK', '负责省ID', '关联网格省表'),
            ('city_id', 'INT', 'FK', '负责市ID', '关联网格市表'),
            ('status', 'TINYINT', '', '状态', '1启用 0禁用'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
    {
        'name': 'supervisor_feedback',
        'cn': '空气质量公众监督反馈表',
        'desc': '存储公众监督员提交的AQI监督反馈信息',
        'fields': [
            ('id', 'INT', 'PK', '反馈ID', '自增主键'),
            ('supervisor_id', 'INT', 'FK', '监督员ID', '关联公众监督员表'),
            ('province_id', 'INT', 'FK', '省ID', '关联网格省表'),
            ('city_id', 'INT', 'FK', '市ID', '关联网格市表'),
            ('grid_address', 'VARCHAR(200)', '', '网格地址', '具体网格位置描述'),
            ('estimated_level', 'INT', 'FK', '预估AQI级别ID', '关联AQI级别表'),
            ('feedback_desc', 'TEXT', '', '反馈描述', '监督员反馈的文字描述'),
            ('feedback_images', 'VARCHAR(500)', '', '反馈图片', '图片URL，逗号分隔'),
            ('status', 'TINYINT', '', '状态', '0待指派 1已指派 2已确认'),
            ('grid_member_id', 'INT', 'FK', '指派网格员ID', '关联网格员表'),
            ('assign_time', 'DATETIME', '', '指派时间', '管理员指派时间'),
            ('confirm_aqi', 'INT', '', '确认AQI值', '网格员确认的实测AQI'),
            ('confirm_level', 'INT', 'FK', '确认AQI级别ID', '关联AQI级别表'),
            ('confirm_time', 'DATETIME', '', '确认时间', '网格员确认时间'),
            ('create_time', 'DATETIME', '', '创建时间', '反馈提交时间'),
        ],
    },
    {
        'name': 'aqi_statistics',
        'cn': '空气质量监测数据统计表',
        'desc': '存储按区域和日期汇总的AQI监测统计数据',
        'fields': [
            ('id', 'INT', 'PK', '统计ID', '自增主键'),
            ('province_id', 'INT', 'FK', '省ID', '关联网格省表'),
            ('city_id', 'INT', 'FK', '市ID', '关联网格市表'),
            ('stat_date', 'DATE', '', '统计日期', '数据统计日期'),
            ('avg_aqi', 'DECIMAL(10,2)', '', '平均AQI', '当日平均AQI值'),
            ('max_aqi', 'INT', '', '最高AQI', '当日最高AQI值'),
            ('min_aqi', 'INT', '', '最低AQI', '当日最低AQI值'),
            ('feedback_count', 'INT', '', '反馈数量', '当日反馈总数'),
            ('confirm_count', 'INT', '', '确认数量', '当日确认总数'),
            ('create_time', 'DATETIME', '', '创建时间', '记录创建时间'),
        ],
    },
]

APIS = [
    {
        'name': '管理员登录',
        'desc': '系统管理员通过账号密码登录，获取访问令牌和管理员信息。',
        'url': '/api/auth/admin/login',
        'method': 'POST',
        'params': [
            ('adminCode', 'String', '是', '管理员账号'),
            ('password', 'String', '是', '登录密码'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": {
    "token": "a1b2c3d4e5f6...",
    "adminInfo": {
      "id": 1,
      "adminCode": "admin",
      "adminName": "系统管理员"
    }
  }
}''',
    },
    {
        'name': '公众监督员注册',
        'desc': '公众监督员注册新账号，填写账号、姓名、密码等信息。',
        'url': '/api/auth/supervisor/register',
        'method': 'POST',
        'params': [
            ('supervisorCode', 'String', '是', '监督员账号'),
            ('supervisorName', 'String', '是', '监督员姓名'),
            ('password', 'String', '是', '登录密码'),
            ('phone', 'String', '否', '联系电话'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": {
    "id": 2,
    "supervisorCode": "user001",
    "supervisorName": "新用户"
  }
}''',
    },
    {
        'name': '公众监督员登录',
        'desc': '公众监督员通过账号密码登录系统。',
        'url': '/api/auth/supervisor/login',
        'method': 'POST',
        'params': [
            ('supervisorCode', 'String', '是', '监督员账号'),
            ('password', 'String', '是', '登录密码'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": {
    "token": "x1y2z3...",
    "supervisorInfo": {
      "id": 1,
      "supervisorCode": "supervisor1",
      "supervisorName": "张三"
    }
  }
}''',
    },
    {
        'name': '公众监督员数据提交',
        'desc': '公众监督员选择网格地址后，提交AQI监督反馈信息，包括预估级别和反馈描述。',
        'url': '/api/supervisor/feedback',
        'method': 'POST',
        'params': [
            ('supervisorId', 'Integer', '是', '监督员ID'),
            ('provinceId', 'Integer', '是', '省ID'),
            ('cityId', 'Integer', '是', '市ID'),
            ('gridAddress', 'String', '是', '网格地址'),
            ('estimatedLevel', 'Integer', '否', '预估AQI级别ID'),
            ('feedbackDesc', 'String', '否', '反馈描述'),
            ('feedbackImages', 'String', '否', '反馈图片URL'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": {
    "id": 1,
    "supervisorId": 1,
    "provinceId": 1,
    "cityId": 1,
    "gridAddress": "沈阳市和平区南京街网格",
    "estimatedLevel": 3,
    "feedbackDesc": "该区域近期空气质量较差",
    "status": 0
  }
}''',
    },
    {
        'name': '管理员查询反馈列表',
        'desc': '系统管理员查询所有公众监督员提交的AQI反馈信息列表。',
        'url': '/api/admin/feedback/list',
        'method': 'GET',
        'params': [],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": [
    {
      "id": 1,
      "supervisorName": "张三",
      "provinceName": "辽宁省",
      "cityName": "沈阳市",
      "gridAddress": "沈阳市和平区南京街网格",
      "levelName": "轻度污染",
      "status": 0,
      "createTime": "2026-06-08T10:30:00"
    }
  ]
}''',
    },
    {
        'name': '管理员指派网格员',
        'desc': '系统管理员将待指派的反馈信息分配给指定网格员进行实地检测。',
        'url': '/api/admin/feedback/assign',
        'method': 'POST',
        'params': [
            ('feedbackId', 'Integer', '是', '反馈ID'),
            ('gridMemberId', 'Integer', '是', '网格员ID'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": "指派成功"
}''',
    },
    {
        'name': '网格员查询指派列表',
        'desc': '网格员查询管理员指派给自己的AQI监督反馈任务列表。',
        'url': '/api/grid/feedback/list',
        'method': 'GET',
        'params': [
            ('gridMemberId', 'Integer', '是', '网格员ID'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": [
    {
      "id": 1,
      "gridAddress": "沈阳市和平区南京街网格",
      "levelName": "轻度污染",
      "feedbackDesc": "该区域近期空气质量较差",
      "status": 1
    }
  ]
}''',
    },
    {
        'name': '网格员确认AQI数据',
        'desc': '网格员实地检测后，提交实测AQI值和确认级别。',
        'url': '/api/grid/feedback/confirm',
        'method': 'POST',
        'params': [
            ('feedbackId', 'Integer', '是', '反馈ID'),
            ('confirmAqi', 'Integer', '是', '实测AQI值'),
            ('confirmLevel', 'Integer', '是', '确认AQI级别ID'),
        ],
        'response': '''{
  "code": 200,
  "message": "success",
  "data": "确认成功"
}''',
    },
]


def set_doc_style(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_table_from_fields(doc, fields, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in fields:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph('')


def add_image_if_exists(doc, image_path, width=Inches(5.5)):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        doc.add_paragraph('')
    else:
        add_para(doc, f'[图片待生成: {os.path.basename(image_path)}]')


def generate_db_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, '东软环保公众监督系统\n数据库设计说明书')

    add_heading(doc, '1. 概述')
    add_para(doc, '本文档描述东软环保公众监督系统的数据库设计方案，包括ER图、数据库设计总表及各数据表的详细字段设计。')
    add_para(doc, '数据库名称：nepm_db')
    add_para(doc, '数据库类型：MySQL 8.0')
    add_para(doc, '字符集：utf8mb4')

    add_heading(doc, '2. ER图')
    er_path = os.path.join(DOCS_DIR, 'er-diagram.png')
    add_image_if_exists(doc, er_path)

    add_heading(doc, '3. 数据库设计总表')
    summary = [(t['name'], t['cn'], t['desc']) for t in TABLES]
    add_table_from_fields(doc, summary, ['表名', '中文名称', '说明'])

    add_heading(doc, '4. 数据表详细设计')
    for t in TABLES:
        add_heading(doc, f'4.{TABLES.index(t)+1} {t["cn"]}（{t["name"]}）', level=2)
        add_para(doc, t['desc'])
        rows = [(f[0], f[1], f[2], f[3], f[4]) for f in t['fields']]
        add_table_from_fields(doc, rows, ['字段名', '类型', '键', '说明', '备注'])

    out = os.path.join(DOCS_DIR, '04-数据库设计说明书.docx')
    doc.save(out)
    print('Generated:', out)


def generate_api_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, '东软环保公众监督系统\n接口设计说明书')

    add_heading(doc, '1. 概述')
    add_para(doc, '本文档描述东软环保公众监督系统主要功能模块的RESTful API接口设计。')
    add_para(doc, '基础URL：http://localhost:8080/api')
    add_para(doc, '统一响应格式：{"code": 200, "message": "success", "data": {...}}')
    add_para(doc, '错误响应格式：{"code": 错误码, "message": "错误信息", "data": null}')

    add_heading(doc, '2. 接口列表')
    for i, api in enumerate(APIS, 1):
        add_heading(doc, f'2.{i} {api["name"]}', level=2)
        add_para(doc, f'功能描述：{api["desc"]}')
        add_para(doc, f'接口地址：{api["url"]}')
        add_para(doc, f'请求方式：{api["method"]}')

        add_para(doc, '请求参数：')
        if api['params']:
            rows = [(p[0], p[1], p[2], p[3]) for p in api['params']]
            add_table_from_fields(doc, rows, ['参数名', '类型', '必填', '说明'])
        else:
            add_para(doc, '无请求参数')

        add_para(doc, '响应示例：')
        p = doc.add_paragraph()
        run = p.add_run(api['response'])
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        doc.add_paragraph('')

    out = os.path.join(DOCS_DIR, '04-接口设计说明书.docx')
    doc.save(out)
    print('Generated:', out)


def generate_req_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, '东软环保公众监督系统\n需求规格说明书')

    add_heading(doc, '1. 项目背景')
    add_para(doc, '随着环境污染问题日益受到社会关注，公众参与环保监督成为提升环境治理效率的重要手段。东软环保公众监督系统旨在构建一个网格化的空气质量（AQI）公众监督平台，通过公众监督员、网格员、系统管理员和决策者四类角色的协同工作，形成"反馈—指派—检测—统计—决策"的闭环监督流程。')

    add_heading(doc, '2. 建设目标')
    add_para(doc, '（1）建立公众参与环保监督的信息化渠道，鼓励公众主动反馈空气质量问题。')
    add_para(doc, '（2）实现网格化管理，将监督反馈精准指派至对应区域的网格员进行实地检测。')
    add_para(doc, '（3）形成AQI监测数据的采集、确认、统计和可视化展示闭环。')
    add_para(doc, '（4）为决策者提供基于实测数据的统计分析和大屏可视化支持。')

    add_heading(doc, '3. 功能需求')
    add_heading(doc, '3.1 公众监督员', level=2)
    add_para(doc, '注册与登录；选择网格地址（省/市/网格）；提交AQI监督反馈（预估级别、描述、图片）。')
    add_heading(doc, '3.2 系统管理员', level=2)
    add_para(doc, '登录系统；查询公众监督员反馈列表；查看反馈详情；指派网格员；查询AQI实测数据。')
    add_heading(doc, '3.3 网格员', level=2)
    add_para(doc, '登录系统；查询指派给自己的监督信息列表；实地检测并确认指派网格的AQI数据。')
    add_heading(doc, '3.4 决策者', level=2)
    add_para(doc, '查询AQI实测数据；生成统计数据；根据统计数据生成可视化大屏。')

    add_heading(doc, '4. 非功能需求')
    add_para(doc, '（1）性能：系统应支持100个并发用户同时在线操作。')
    add_para(doc, '（2）安全：采用Token鉴权机制，敏感接口需验证登录状态。')
    add_para(doc, '（3）可用性：界面简洁友好，移动端适配公众监督员和网格员使用场景。')
    add_para(doc, '（4）可维护性：前后端分离架构，接口规范统一，便于扩展。')

    add_heading(doc, '5. 软硬件环境')
    add_para(doc, '客户端：Chrome/Firefox/Edge 现代浏览器；移动端浏览器。')
    add_para(doc, '服务端：JDK 8+、Spring Boot 2.7、MyBatis、MySQL 8.0。')
    add_para(doc, '前端：Vue 3、Vite、Element Plus、Pinia、Vue Router、Axios。')
    add_para(doc, '开发工具：IntelliJ IDEA、VS Code/Cursor、Maven、Node.js。')

    add_heading(doc, '6. 模块细分')
    modules = [
        ('认证模块', '管理员/监督员/网格员登录注册'),
        ('反馈模块', '公众监督员提交AQI监督反馈'),
        ('指派模块', '管理员查询反馈并指派网格员'),
        ('检测确认模块', '网格员实地检测并确认AQI数据'),
        ('统计模块', 'AQI数据统计与汇总'),
        ('可视化模块', '决策者大屏展示'),
    ]
    add_table_from_fields(doc, modules, ['模块名称', '功能说明'])

    add_heading(doc, '7. 业务流程')
    add_para(doc, '系统主业务流程如下：公众监督员注册登录后选择网格并提交反馈；管理员查询反馈并指派网格员；网格员查询指派信息并实地检测确认AQI；决策者基于实测数据生成统计和可视化大屏。')
    flow_images = ['主业务流程图.png', '公众反馈流程图.png', '网格员确认流程图.png', '管理员指派流程图.png', '统计看板流程图.png']
    for img in flow_images:
        add_image_if_exists(doc, os.path.join(DOCS_DIR, '03-流程图', img), width=Inches(5))

    add_heading(doc, '8. 原型概述')
    add_para(doc, '系统包含三个终端：系统管理员端（Web后台）、公众监督员端（移动端）、网格员端（移动端）。')
    proto_images = [
        ('系统管理员端-登录.png', '管理员登录页面'),
        ('系统管理员端-反馈列表.png', '反馈列表页面'),
        ('公众监督员端-登录.png', '监督员登录页面'),
        ('公众监督员端-提交反馈.png', '提交反馈页面'),
        ('网格员端-指派列表.png', '网格员指派列表'),
        ('网格员端-确认AQI数据.png', '网格员确认AQI页面'),
    ]
    for img, desc in proto_images:
        add_para(doc, desc + '：')
        add_image_if_exists(doc, os.path.join(DOCS_DIR, '03-原型图', img), width=Inches(3))

    out = os.path.join(DOCS_DIR, '05-需求规格说明书.docx')
    doc.save(out)
    print('Generated:', out)


if __name__ == '__main__':
    os.makedirs(DOCS_DIR, exist_ok=True)
    generate_db_doc()
    generate_api_doc()
    generate_req_doc()
